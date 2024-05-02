import streamlit as st
import pyttsx3
import docx
import webbrowser
import os
import pyautogui
import cv2
import pygame
import sys
import nltk
import spacy
import numpy as np
import requests
import speech_recognition as sr
import sounddevice as sd
import soundfile as sf
import seaborn as sns
import plotly.express as px
import pandas as pd
import matplotlib.pyplot as plt
from bs4 import BeautifulSoup
from PyPDF2 import PdfReader
from statsmodels.tsa.seasonal import seasonal_decompose
from sklearn.linear_model import LinearRegression
from textblob import TextBlob
from heapq import nlargest
from datetime import datetime

st.set_page_config(
    page_title="Asistent Vocal",
    page_icon="👋",
)

# Funcția pentru încărcarea datelor din fișier
def load_data(uploaded_file):
    if uploaded_file is not None:
        try:
            data = pd.read_excel(uploaded_file, index_col="ds")
            return data
        except Exception as e:
            st.error(f"A apărut o eroare la încărcarea fișierului: {str(e)}")
    else:
        return None

def show_analysis(data):
    if data is not None:
        st.subheader("Datele despre calitatea aerului")
        st.write(data)

        # EDA
        st.markdown("### 🕵️‍♂️ Analiza datelor:", unsafe_allow_html=True)
        with st.expander("Afișați EDA"):
            if st.checkbox("Afișați primele 5 rînduri"):
                st.write(data.head())
            if st.checkbox("Afișați ultimele 5 rînduri"):
                st.write(data.tail())
            if st.checkbox("Afișare formă"):
                st.write(f"Acest DataFrame are **{data.shape[0]} rinduri** și **{data.shape[1]} coloane**.")
            if st.checkbox("Afișare coloane"):
                st.write(pd.DataFrame(data.columns, columns=['Columns']).T)
            if st.checkbox("Afișați descrierea **(Caractersitici numerice)**"):
                st.write(data.describe())
            if st.checkbox("Afișați descrierea **(Caractersitici categorice)**"):
                if data.select_dtypes(include=object).columns.tolist():
                    st.write(data.describe(include=['object']))
                else:
                    st.info("Nu există caracteristici categorice.")
            if st.checkbox("Afișare corelație"):
                if data.corr().columns.tolist():
                    fig, ax = plt.subplots()
                    sns.heatmap(data.corr(), cmap='Blues', annot=True, ax=ax)
                    st.pyplot(fig)
                else:
                    st.info("Nu există caracteristici de corelație.")
            if st.checkbox("Afișare valori lipsă"):
                col1, col2 = st.columns([0.4, 1])
                with col1:
                    st.markdown("<h6 align='center'> Numar de valori nule", unsafe_allow_html=True)
                    st.write(data.isnull().sum().sort_values(ascending=False))
                with col2:
                    st.markdown("<h6 align='center'> Reprezentare grafică pentru valorile nule ", unsafe_allow_html=True)
                    null_values = data.isnull().sum()
                    null_values = null_values[null_values > 0]
                    if not null_values.empty:
                        null_values = null_values.sort_values(ascending=False).to_frame()
                        null_values.columns = ['Count']
                        null_values.index.names = ['Feature']
                        null_values['Feature'] = null_values.index
                        fig = px.bar(null_values, x='Feature', y='Count', color='Count')
                        st.plotly_chart(fig)
                    else:
                        st.info("Nu sunt valori lipsă.")

            if st.checkbox("Ștergerea coloanelor"):
                col_to_delete = st.multiselect("Selectați coloanele de șters", data.columns)
                if st.button("Șterge"):
                    data.drop(columns=col_to_delete, inplace=True)
                    st.success(f"Coloana  **`{col_to_delete}`** a fost ștearsă cu succes")

            if st.button("Afișare DataFrame"):
                st.write(data)

        # Analiză de serie temporală
        st.markdown("### 📈 Analiză de Serie Temporală", unsafe_allow_html=True)
        if isinstance(data.index, pd.DatetimeIndex):
            time_series_var = st.selectbox("Selectați Variabila:", data.columns)
            fig = px.line(data, x=data.index, y=time_series_var,
                          title=f"Evoluția lui {time_series_var} în timp")
            st.plotly_chart(fig)
            st.subheader("Analiza Sezonierității")
            seasonal_decomposition = seasonal_decompose(data[time_series_var], model='additive')
            fig, (ax1, ax2, ax3, ax4) = plt.subplots(4, 1, figsize=(10, 8))
            seasonal_decomposition.observed.plot(ax=ax1)
            ax1.set_ylabel('Observat')
            seasonal_decomposition.trend.plot(ax=ax2)
            ax2.set_ylabel('Tendință')
            seasonal_decomposition.seasonal.plot(ax=ax3)
            ax3.set_ylabel('Sezonier')
            seasonal_decomposition.resid.plot(ax=ax4)
            ax4.set_ylabel('Residual')
            plt.tight_layout()
            st.pyplot(fig)

            # Adăugarea analizei tendințelor temporale
            st.subheader("Analiza Tendințelor Temporale")
            df = data[[time_series_var]].copy()
            df.dropna(inplace=True)
            df['Time'] = np.arange(len(df))
            X = df[['Time']]
            y = df[time_series_var]
            model = LinearRegression()
            model.fit(X, y)
            st.write(f"Tendința liniară a variabilei '{time_series_var}': {model.coef_[0]:.2f} unități pe unitate de timp")
            plt.figure(figsize=(10, 6))
            plt.plot(df.index, df[time_series_var], label='Actual Data')
            plt.plot(df.index, model.predict(X), label='Trend Line', linestyle='--', color='red')
            plt.title(f"Tendința Temporală a Variabilei '{time_series_var}'")
            plt.xlabel("Timp")
            plt.ylabel("Valoare")
            plt.legend()
            st.pyplot()

        else:
            st.warning("DataFrame-ul nu este indexat cu date timp. Vă rugăm să specificați coloana de date timp.")

def main():
    st.title("Analiză Date")
    uploaded_file = st.file_uploader("Încărcați fișierul DataFrame (.csv, .xlsx, .txt)", type=["csv", "xlsx", "txt"])
    if uploaded_file:
        data = load_data(uploaded_file)
        show_analysis(data)
    else:
        st.info("Vă rugăm să încărcați un fișier pentru a începe analiza.")

########################################################################################
# Funcția pentru a prelua comanda vocală utilizând SpeechRecognition
def start_listening():
    recognizer = sr.Recognizer()
    microphone = sr.Microphone()

    while True:
        with microphone as source:
            recognizer.adjust_for_ambient_noise(source)
            with st.spinner("Ascult..."):
                audio = recognizer.listen(source)
        try:
            with st.spinner("Procesez..."):
                command = recognizer.recognize_google(audio, language="ro-RO").lower()
                st.success(f"Ai spus: {command}")  # Afiseaza comanda sub forma de cadru de succes
                return command
        except sr.UnknownValueError:
            st.error("Scuze, nu am înțeles. Te rog să repeți comanda.")  # Afiseaza o eroare in cazul in care comanda nu este recunoscuta
        except sr.RequestError:
            st.error("Scuze, a apărut o eroare în procesarea cererii tale. Te rog să repeți comanda.")  # Afiseaza o eroare in cazul unei erori in procesare

AUDIO_DIRECTORY = "audio"

if not os.path.exists(AUDIO_DIRECTORY):
    os.makedirs(AUDIO_DIRECTORY)

def record_voice_note():
    st.write("Cum doriți să numiți fișierul?")
    speak("Cum doriți să numiți fișierul?")
    file_name = start_listening()
    speak(f"Începeți înregistrarea pentru fișierul {file_name}. Vorbiți acum...")
    st.write(f"Începeți înregistrarea pentru fișierul {file_name}. Vorbiți acum...")
    duration = 10
    recording = sd.rec(int(duration * 44100), samplerate=44100, channels=2, dtype='int16')
    sd.wait()
    file_path = os.path.join(AUDIO_DIRECTORY, f"{file_name}.wav")
    sf.write(file_path, recording, 44100)
    st.write("Înregistrare finalizată.")
    speak("Înregistrare finalizată.")

def play_voice_note():
    st.write("Despre ce fișier doriți să ascultați?")
    speak("Despre ce fișier doriți să ascultați?")
    file_name = start_listening()
    st.write(f"Redare fișier {file_name}...")
    speak(f"Redare fișier {file_name}...")
    file_path = os.path.join(AUDIO_DIRECTORY, f"{file_name}.wav")
    data, samplerate = sf.read(file_path)
    sd.play(data, samplerate)
    sd.wait()
    sentiment, summary, entities = analyze_and_summarize_audio(file_path)  # Modificarea aici
    display_results(sentiment, summary, entities)  # Modificarea aici

def speak(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()

def transcribe_audio(file_path):
    recognizer = sr.Recognizer()
    audio_file = sr.AudioFile(file_path)
    with audio_file as source:
        audio_data = recognizer.record(source)
    try:
        transcript = recognizer.recognize_google(audio_data, language="ro-RO")
        return transcript
    except sr.UnknownValueError:
        print("Recunoașterea vocală nu a putut interpreta audio.")
        return ""
    except sr.RequestError:
        print("Eroare de conexiune la serviciul de recunoaștere vocală.")
        return ""

def analyze_sentiment(text):
    blob = TextBlob(text)
    polarity = blob.sentiment.polarity
    return polarity

def analyze_sentiment_category(polarity):
    if polarity > 0:
        return "Pozitiv"
    elif polarity < 0:
        return "Negativ"
    else:
        return "Neutru"

def summarize_text(text):
    sentences = nltk.sent_tokenize(text)
    word_frequencies = {}
    for sentence in sentences:
        for word in nltk.word_tokenize(sentence):
            if word not in word_frequencies:
                word_frequencies[word] = 1
            else:
                word_frequencies[word] += 1
    maximum_frequency = max(word_frequencies.values())
    for word in word_frequencies.keys():
        word_frequencies[word] = (word_frequencies[word] / maximum_frequency)
    sentence_scores = {}
    for sentence in sentences:
        for word in nltk.word_tokenize(sentence.lower()):
            if word in word_frequencies:
                if len(sentence.split(' ')) < 30:
                    if sentence not in sentence_scores:
                        sentence_scores[sentence] = word_frequencies[word]
                    else:
                        sentence_scores[sentence] += word_frequencies[word]
    summary_sentences = nlargest(7, sentence_scores, key=sentence_scores.get)
    summary = ' '.join(summary_sentences)
    return summary

# Încărcarea modelului SpaCy pentru limba română
nlp_ro = spacy.load("ro_core_news_sm")

def detect_entities(text):
    doc = nlp_ro(text)
    entities = [(ent.text, ent.label_) for ent in doc.ents]
    return entities


def analyze_and_summarize_audio(file_path):
    transcript = transcribe_audio(file_path)
    entities = detect_entities(transcript)
    polarity = analyze_sentiment(transcript)
    sentiment_category = analyze_sentiment_category(polarity)
    summary = summarize_text(transcript)
    return sentiment_category, summary, entities

def display_results(sentiment, summary, entities):
    st.write("Sentiment:", sentiment)
    st.write("Rezumat:", summary)
    st.write("Entități detectate:")
    for entity in entities:
        st.write(entity)

def open_web_browser():
    speak("Deschid browser-ul web. Ce dorești să cauți sau să accesezi?")
    while True:
        command = start_listening()
        if "caută" in command:
            speak("Ce dorești să cauți pe internet?")
            query = start_listening()
            speak(f"Caut pe internet pentru: {query}")

            search_url = f'https://www.google.com/search?q={query}'
            response = requests.get(search_url)
            soup = BeautifulSoup(response.text, 'html.parser')

            search_results = soup.find_all('a')
            first_result = None
            for link in search_results:
                href = link.get('href')
                if href.startswith('/url?q='):
                    first_result = href[7:]
                    break

            if first_result:
                st.write("Accesam pagina.")
                speak("Accesam pagina.")
                webbrowser.open_new_tab(first_result)
            else:
                st.write("Nu am putut găsi niciun rezultat pentru căutarea ta.")
                speak("Nu am putut găsi niciun rezultat pentru căutarea ta.")
        elif "deschide" in command:
            st.write("Despre ce pagină web este vorba?")
            speak("Despre ce pagină web este vorba?")
            page_name = start_listening()
            speak(f"Accesez pagina {page_name}.")
            webbrowser.open_new_tab(f"http://{page_name}.com")

            search_results = soup.find_all('a')
            first_result = None
            for link in search_results:
                href = link.get('href')
                if href.startswith('/url?q='):
                    first_result = href[7:]
                    break

            if first_result:
                st.write("Accesând pagina.")
                speak("Accesând pagina.")
                webbrowser.open_new_tab(first_result)
            else:
                st.write("Nu am putut găsi niciun rezultat pentru căutarea ta.")
                speak("Nu am putut găsi niciun rezultat pentru căutarea ta.")
        elif "ieșire" in command:
            st.write("Închid browser-ul web.")
            speak("Închid browser-ul web.")
            break

        else:
            speak("Comandă invalidă. Te rog să spui 'caută', 'deschide' sau 'ieșire'.")


net = cv2.dnn.readNet("yolov3.weights", "yolov3.cfg")
classes = []
with open("coco.names", "r") as f:
    classes = [line.strip() for line in f.readlines()]
layer_names = net.getLayerNames()
output_layers = [layer_names[i - 1] for i in net.getUnconnectedOutLayers()]

detected_objects_prev = set()

def start_camera_with_voice_command():
    r = sr.Recognizer()
    with sr.Microphone() as source:
        st.write("Ziceți „start” pentru a porni camera sau „stop” pentru a opri camera...")
        r.adjust_for_ambient_noise(source)
        audio = r.listen(source)

        try:
            text = r.recognize_google(audio, language="ro-RO")
            if "start" in text.lower():
                start_camera()
            elif "stop" in text.lower():
                st.write("Camera oprită.")
            else:
                st.write("Comanda vocală nu a fost recunoscută sau nu a fost pornită/oprită camera.")
        except sr.UnknownValueError:
            st.write("Nu am putut recunoaște comanda vocală.")
        except sr.RequestError:
            st.write("Nu am putut accesa serviciul de recunoaștere vocală. Verificați conexiunea la internet.")

def start_camera():
    global detected_objects_prev
    camera_capture = cv2.VideoCapture(0)
    display_interval = 100
    frame_count = 0
    while True:
        is_true, frame = camera_capture.read()
        height, width, channels = frame.shape
        frame_count += 1
        if frame_count % display_interval == 0:
        # Detectează obiectele folosind YOLO
            blob = cv2.dnn.blobFromImage(frame, 0.00392, (416, 416), (0, 0, 0), True, crop=False)
            net.setInput(blob)
            outs = net.forward(output_layers)
            detected_objects_curr = set()
            for out in outs:
                for detection in out:
                    scores = detection[5:]
                    class_id = np.argmax(scores)
                    confidence = scores[class_id]
                    if confidence > 0.5:
                        detected_object = classes[class_id]
                        detected_objects_curr.add(detected_object)
                        center_x = int(detection[0] * width)
                        center_y = int(detection[1] * height)
                        w = int(detection[2] * width)
                        h = int(detection[3] * height)
                        x = int(center_x - w / 2)
                        y = int(center_y - h / 2)
                        cv2.rectangle(frame, (x, y), (x + w, y + h), (0, 255, 0), 2)
                        cv2.putText(frame, detected_object, (x, y + 30), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
            new_objects = detected_objects_curr - detected_objects_prev
            st.image(frame, channels="BGR")
            st.write(list(new_objects))
            detected_objects_prev = detected_objects_curr
        if cv2.waitKey(1) & 0xFF == ord('q'):
            break
    camera_capture.release()
    cv2.destroyAllWindows()

def take_screenshot():
    st.write("Fac un screenshot.")
    speak("Fac un screenshot.")
    screenshot = pyautogui.screenshot()
    st.image(screenshot, caption='Screenshot', use_column_width=True)

def take_picture():
    st.write("Fac o fotografie.")
    speak("Fac o fotografie.")
    camera = cv2.VideoCapture(0)
    return_value, image = camera.read()
    cv2.imwrite('photo.png', image)
    camera.release()
    st.image('photo.png', caption='Photography', use_column_width=True)

def open_presentation():
    st.write("Deschid prezentarea.")
    speak("Deschid prezentarea.")
    presentation_path = "prezentare.pptx"
    os.startfile(presentation_path)

def read_word_file(filename):
    doc_exists = os.path.exists(filename + '.docx')
    if not doc_exists:
        doc = docx.Document()
        doc.add_paragraph("Acesta este un fișier Word creat prin comandă vocală.")
        doc.save(filename + '.docx')
        speak(f"Fișierul Word '{filename}.docx' a fost creat cu succes.")

    doc = docx.Document(filename + '.docx')
    speak("Ce dorești să faci cu fișierul Word?")
    action = start_listening()

    if action == "adaugare" or action == "adăugare":
        speak("Spuneți textul pe care doriți să-l adăugați.")
        text_to_add = start_listening()
        doc.add_paragraph(text_to_add)
        doc.save(filename + '.docx')
        speak("Textul a fost adăugat cu succes în fișierul Word.")

    elif action == "ștergere":
        speak("Spuneți textul pe care doriți să-l ștergeți.")
        text_to_delete = start_listening()
        paragraphs_to_delete = [p for p in doc.paragraphs if text_to_delete in p.text]
        for p in paragraphs_to_delete:
            p.clear()
        doc.save(filename + '.docx')
        speak("Textul a fost șters cu succes din fișierul Word.")

    elif action == "citire":
        speak("Citește fișierul Word.")
        for paragraph in doc.paragraphs:
            speak(paragraph.text)

def read_pdf_file(filename):
    speak("Ce dorești să faci cu fișierul PDF?")
    action = st.radio("Alege acțiunea:", ("Citire", "Adăugare", "Ștergere"))

    if action == "Adăugare":
        text_to_add = st.text_input("Introdu textul pe care dorești să-l adaugi:")
        if st.button("Adaugă text"):
            with open(filename + '.pdf', 'rb') as file:
                pdf_reader = PdfReader(file)
                pdf_writer = PdfWriter()
                for page_number, page in enumerate(pdf_reader.pages):
                    page.text_objects.insert(0, create_text_object(text_to_add))
                    pdf_writer.add_page(page)
                with open('new_' + filename + '.pdf', 'wb') as new_file:
                    pdf_writer.write(new_file)
            speak("Textul a fost adăugat cu succes în fișierul PDF.")

    elif action == "Ștergere":
        text_to_delete = st.text_input("Introdu textul pe care dorești să-l ștergi:")
        if st.button("Șterge text"):
            with open(filename + '.pdf', 'rb') as file:
                pdf_reader = PdfReader(file)
                pdf_writer = PdfWriter()
                for page_number, page in enumerate(pdf_reader.pages):
                    page.text_objects = [to for to in page.text_objects if text_to_delete not in to.get_text()]
                    pdf_writer.add_page(page)
                with open('new_' + filename + '.pdf', 'wb') as new_file:
                    pdf_writer.write(new_file)
            speak("Textul a fost șters cu succes din fișierul PDF.")

    elif action == "Citire":
        speak("Citește fișierul PDF.")
        with open(filename + '.pdf', 'rb') as file:
            pdf_reader = PdfReader(file)
            for page_number, page in enumerate(pdf_reader.pages):
                st.write(page.extract_text())

def create_text_object(text):
    text_object = TextObject()
    text_object.set_text(text)
    return text_object

def get_weather(city):
    api_key = "Your API key "  # Înlocuiți cu cheia dvs. API OpenWeatherMap
    url = f"http://api.openweathermap.org/data/2.5/weather?q={city}&appid={api_key}&units=metric"
    response = requests.get(url)
    data = response.json()
    temperature = data['main']['temp']
    humidity = data['main']['humidity']
    precipitation = data['weather'][0]['main']
    return temperature, humidity, precipitation

def read_text(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()

def main_control_vocal():
    st.title("Control vocal")
    st.sidebar.title("Meniu cu control vocal")
    speak("Spune 'Începe' pentru a începe controlul vocal.")
    menu_options = [
        "Pornire vocală",
        "Inregistrare vocală(inregistrare)",
        "Redare vocală(redare)",
        "Deschide browser-ul web(browser)",
        "Deschide o prezentare(prezentare)",
        "Detectează obiecte(camera)",
        "Fă o fotografie(poza)",
        "Fă un screenshot(screenshot)",
        "Citește un fișier Word(word)",
        "Citește un fișier PDF(pdf)",
        "Afișează temperatura(temperatura)",
        "Ieșire(iesire)"
    ]
    selected_option = st.sidebar.radio("Alege o opțiune:", menu_options)

    camera_started = False

    if selected_option == "Pornire vocală":
        st.write("Spune 'Pornire' pentru a începe controlul vocal.")
        speak("Spune 'Pornire' pentru a începe controlul vocal.")
        command = start_listening()
        if command == "pornire":
            st.write("Controlul vocal a fost pornit. Spune o comandă din meniu.")

            while True:
                command = start_listening()
                if command == "joc":
                    start_game()
                elif command == "înregistrare":
                    record_voice_note()
                elif command == "redare":
                    play_voice_note()
                elif command == "browser":
                    open_web_browser()
                elif command == "prezentare":
                    open_presentation()
                elif command == "cameră" or command =="camera":
                    if not camera_started:
                        start_camera()
                        camera_started = True
                elif command == "oprire":
                    if camera_started:
                        st.write("Camera oprită.")
                        break
                elif command == "poză":
                    take_picture()
                elif command == "screenshot":
                    take_screenshot()
                elif command == "document pdf" or command == "pdf":
                    speak("Spune numele fișierului PDF pe care dorești să-l deschizi.")
                    filename = start_listening()
                    if filename:
                        read_pdf_file(filename)
                elif command == "document word" or command == "word":
                    speak("Spune numele fișierului word pe care dorești să-l deschizi.")
                    filename = start_listening()
                    if filename:
                        read_word_file(filename)
                elif command == "temperatura" or command == "temperatură":
                    speak("Despre ce oraș dorești să vezi temperatura?")
                    st.write("Despre ce oraș dorești să vezi temperatura?")
                    city = start_listening()
                    try:
                        temperature, humidity, precipitation = get_weather(city)
                        speak(f'În {city}:\nTemperatura este de aproximativ {temperature} grade Celsius.')
                        st.write(f'În {city}:\nTemperatura este de aproximativ {temperature} grade Celsius.')
                        speak(f'Umiditatea este de {humidity} %.')
                        st.write(f'Umiditatea este de {humidity} %.')
                        speak(f'Condiția meteorologică este {precipitation}.')
                        st.write(f'Condiția meteorologică este {precipitation}.')
                    except Exception as e:
                        speak(
                            "Nu am putut obține informațiile meteo. Verificați conexiunea la internet sau încercați mai târziu.")
                elif command == "ieșire":
                    st.write("La revedere")
                    speak("La revedere")
                    st.stop()
                    break

    elif selected_option == "Inregistrare voce":
        record_voice_note()
    elif selected_option == "Redare voce":
        play_voice_note()
    elif selected_option == "Deschide browser-ul web":
        open_web_browser()
    elif selected_option == "Deschide o prezentare":
        open_presentation()
    elif selected_option == "Detectează obiecte":
        start_camera()
    elif selected_option == "Fă o fotografie":
        take_picture()
    elif selected_option == "Fă un screenshot":
        take_screenshot()
    elif selected_option == "Citește un fișier Word":
        filename = st.sidebar.text_input("Introdu numele fișierului Word:")
        read_word_file(filename)
    elif selected_option == "Citește un fișier PDF":
        filename = st.sidebar.text_input("Introdu numele fișierului PDF:")
        read_pdf_file(filename)
    elif selected_option == "Citire text":
        text_to_read = st.text_input("Introduceți textul de citit:")
        if st.button("Citire text"):
            read_text(text_to_read)
    elif selected_option == "Afișează temperatura":
        speak("Despre ce oraș dorești să vezi temperatura?")
        city = start_listening()
        try:
            temperature, humidity, precipitation = get_weather(city)
            speak(f'În {city}:\nTemperatura este de aproximativ {temperature} grade Celsius.')
            st.write(f'În {city}:\nTemperatura este de aproximativ {temperature} grade Celsius.')
            speak(f'Umiditatea este de {humidity} %.')
            st.write(f'Umiditatea este de {humidity} %.')
            speak(f'Condiția meteorologică este {precipitation}.')
            st.write(f'Condiția meteorologică este {precipitation}.')
        except Exception as e:speak("Nu am putut obține informațiile meteo. Verificați conexiunea la internet sau încercați mai târziu.")
    elif selected_option == "Ieșire":
        st.stop()
        speak("La revedere")

def main_incarcare_date():
    st.title("Încărcare Date")
    uploaded_file = st.file_uploader("Încărcați fișierul DataFrame (.csv, .xlsx, .txt)", type=["csv", "xlsx", "txt"])
    if uploaded_file:
        data = load_data(uploaded_file)
        show_analysis(data)

def help_section():
    st.subheader("Ajutor - Comenzi Vocale")
    commands = {
        "Inregistrare voce": "Comanda 'inregistrare' va inregistra vocea care o veti rosti. Va va intreba cum doriti sa denumiti fisierul.",
        "Redare voce": "Comanda 'redare' va reda vocea/fisierul din folderul audio pe care l-ati creat, de asemenea va arata rezumatul, sentimentul si entitatea din inregistrea facuta.",
        "Deschide browser-ul web": "Comanda 'web' va deschide un nou tab în browser-ul web."
                                   "Comanda 'cauta' va deschide o pagina care doriti a cauta pe internet."
                                   "Comanda 'deschide' va deschide o pagina noua .com"
                                    "Comanda 'iesire' va inchide browserul web",
        "Deschide o prezentare": "Comanda 'prezentare' va deschide o prezentare PowerPoint.",
        "Fă o fotografie": "Comanda 'poză' va activa camera și va face o fotografie.",
        "Fă un screenshot": "Comanda 'screenshot' va realiza un screenshot al ecranului.",
        "Citește un fișier Word": "Comanda 'Word' va deschide un document word iar cu comenzile.... "
                                  "'adauga' va adauga text in documentul word, "
                                  "'stergere' va sterge textul din documentul word, "
                                  "'citire' va citi textul din documentul word.",
        "Citește un fișier PDF": "Comanda 'PDF' va citi conținutul unui fișier PDF.",
        "Afișează temperatura": "Comanda 'Temperatura' va afișa temperatura și condițiile meteorologice curente pentru un anumit oraș.",
        "Ieșire": "Comanda 'Ieșire' va închide aplicația."
    }

    for command, description in commands.items():
        st.write(f"- **{command}:** {description}")

def display_current_datetime():
    current_datetime = datetime.now()
    current_time = current_datetime.strftime("%H:%M:%S")
    current_date = current_datetime.strftime("%d-%m-%Y")
    st.sidebar.write(f"Data/Ora: {current_date} {current_time}")

#################################################################
# Funcția principală
def main():
    display_current_datetime()
    st.sidebar.title("Meniu Principal")
    page_options = ["Control Vocal", "Încărcare Date", "Joc Pong Game"]
    selected_page = st.sidebar.radio("Selectează o pagină:", page_options)
    website_heading = "Assistent vocal pentru persoane cu dizabilitati"
    st.markdown(f"<h1 style='text-align: center; color: blue;'>{website_heading}</h1>", unsafe_allow_html=True)
    help_section()

    if selected_page == "Control Vocal":
        main_control_vocal()
    elif selected_page == "Încărcare Date":
        main_incarcare_date()
    elif selected_page == "Joc Pong Game":
        start_game()

#################################################################
def start_game():
    st.write("Pornire joc...")
    speak("Pornire joc...")
    pygame.init()

    WIDTH, HEIGHT = 600, 400
    WHITE = (255, 255, 255)
    BLACK = (0, 0, 0)
    BALL_SPEED = 5

    screen = pygame.display.set_mode((WIDTH, HEIGHT))
    pygame.display.set_caption("Pong Game")

    ball_x, ball_y = WIDTH // 2, HEIGHT // 2
    ball_speed_x, ball_speed_y = BALL_SPEED, BALL_SPEED

    paddle_width, paddle_height = 15, 60
    left_paddle_x, right_paddle_x = 10, WIDTH - 25
    left_paddle_y, right_paddle_y = HEIGHT // 2 - paddle_height // 2, HEIGHT // 2 - paddle_height // 2
    paddle_speed = 7
    score_left, score_right = 0, 0

    font = pygame.font.Font(None, 36)

    def reset_ball():
        return WIDTH // 2, HEIGHT // 2, BALL_SPEED, BALL_SPEED

    while True:
        for event in pygame.event.get():
            if event.type == pygame.QUIT:
                pygame.quit()
                sys.exit()

        keys = pygame.key.get_pressed()
        if keys[pygame.K_w] and left_paddle_y > 0:
            left_paddle_y -= paddle_speed
        if keys[pygame.K_s] and left_paddle_y < HEIGHT - paddle_height:
            left_paddle_y += paddle_speed
        if keys[pygame.K_UP] and right_paddle_y > 0:
            right_paddle_y -= paddle_speed
        if keys[pygame.K_DOWN] and right_paddle_y < HEIGHT - paddle_height:
            right_paddle_y += paddle_speed

        ball_x += ball_speed_x
        ball_y += ball_speed_y

        if (
            left_paddle_x < ball_x < left_paddle_x + paddle_width
            and left_paddle_y < ball_y < left_paddle_y + paddle_height
        ) or (
            right_paddle_x < ball_x < right_paddle_x + paddle_width
            and right_paddle_y < ball_y < right_paddle_y + paddle_height
        ):
            ball_speed_x = -ball_speed_x

        if ball_y <= 0 or ball_y >= HEIGHT:
            ball_speed_y = -ball_speed_y

        if ball_x <= 0:
            score_right += 1
            ball_x, ball_y, ball_speed_x, ball_speed_y = reset_ball()

        if ball_x >= WIDTH:
            score_left += 1
            ball_x, ball_y, ball_speed_x, ball_speed_y = reset_ball()

        screen.fill(BLACK)
        pygame.draw.rect(screen, WHITE, (left_paddle_x, left_paddle_y, paddle_width, paddle_height))
        pygame.draw.rect(screen, WHITE, (right_paddle_x, right_paddle_y, paddle_width, paddle_height))
        pygame.draw.ellipse(screen, WHITE, (ball_x - 10, ball_y - 10, 20, 20))
        score_display = font.render(f"{score_left} - {score_right}", True, WHITE)
        screen.blit(score_display, (WIDTH // 2 - 40, 10))

        pygame.display.flip()
        pygame.time.Clock().tick(60)

if __name__ == "__main__":
    main()
